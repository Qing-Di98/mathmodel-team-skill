# -*- coding: utf-8 -*-
"""提取课件 docx 为结构化 md（标题层级 + 正文 + 目录），供 skill 检索学习。

一步完成：提取全部 docx 后自动运行 add_toc.py 生成"## 目录"
（add_toc.py 幂等，可单独重跑）。课件更新流程：复制新 docx → 运行本脚本。
"""
import os
import re
import subprocess
import sys
from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

BASE = os.path.dirname(os.path.abspath(__file__))
TARGETS = [
    ("CUMCM国赛案例与建模步骤.docx", "CUMCM国赛案例与建模步骤.md"),
    ("建模思维与流程.docx", "建模思维与流程.md"),
    ("数学建模新手备赛指南.docx", "数学建模新手备赛指南.md"),
    ("数学建模竞赛备赛教案.docx", "数学建模竞赛备赛教案.md"),
    ("数模第一节课补充.docx", "数模第一节课补充.md"),
    ("优化模型与评价模型/AHP.docx", "优化模型与评价模型/AHP.md"),
    ("优化模型与评价模型/TOPSIS.docx", "优化模型与评价模型/TOPSIS.md"),
    ("优化模型与评价模型/bigM.docx", "优化模型与评价模型/bigM.md"),
    ("优化模型与评价模型/PuLP代码.docx", "优化模型与评价模型/PuLP代码.md"),
    ("优化模型与评价模型/优化模型 + 评价模型.docx", "优化模型与评价模型/优化模型 + 评价模型.md"),
    ("数据处理与预测模型/1数据预处理理论基础.docx",
     "数据处理与预测模型/1数据预处理理论基础.md"),
    ("数据处理与预测模型/2ARIMA&SARIMA时间序列模型零基础全解.docx",
     "数据处理与预测模型/2ARIMA&SARIMA时间序列模型零基础全解.md"),
    ("数据处理与预测模型/3灰色GM(1,1)完整数学推导、残差检验、封装类代码.docx",
     "数据处理与预测模型/3灰色GM(1,1)完整数学推导、残差检验、封装类代码.md"),
    ("数据处理与预测模型/数模国赛聚类分析（K-Means）.docx",
     "数据处理与预测模型/数模国赛聚类分析（K-Means）.md"),
    ("数据处理与预测模型/GM(1,1)灰色预测模型在农作物年产量预测中的应用——基于2019-2023年数据的实证分析.docx",
     "数据处理与预测模型/GM(1,1)灰色预测模型在农作物年产量预测中的应用——基于2019-2023年数据的实证分析.md"),
    ("数学建模题型与算法对照表.docx", "数学建模题型与算法对照表.md"),
]


def heading_level(p):
    try:
        return int(p.style.name.split()[-1]) if p.style.name.startswith("Heading") else 0
    except Exception:
        return 0


def clean(t):
    t = re.sub(r"\s+", " ", t).strip()
    return t


def cell_text_marked(cell):
    """单元格文本 + 格式标注：绿色高亮 🟢（优先）、删除线 ⛔（尽量不用）。

    对照表类课件用绿色高亮/划横线表达模型选型偏好，提取为 md 时保留标注。
    """
    parts = []
    for p in cell.paragraphs:
        seg = []
        for r in p.runs:
            t = r.text
            if t == "":
                continue
            rPr = r._element.rPr
            marks = ""
            if rPr is not None:
                hl = rPr.find(qn("w:highlight"))
                if hl is not None and hl.get(qn("w:val")) == "green":
                    marks += "🟢"
                if rPr.find(qn("w:strike")) is not None:
                    marks += "⛔"
            seg.append(t + marks)
        parts.append("".join(seg))
    return clean("\n".join(parts))


def run_add_toc():
    """提取后自动生成目录（复用 add_toc.py，幂等；失败仅告警不中断）。"""
    add_toc = os.path.join(BASE, "add_toc.py")
    r = subprocess.run([sys.executable, add_toc],
                       capture_output=True, text=True, encoding="utf-8")
    sys.stdout.write(r.stdout)
    if r.stderr:
        sys.stderr.write(r.stderr)
    if r.returncode != 0:
        print(f"WARN: add_toc.py 退出码 {r.returncode}，请手动重跑")


def main():
    for src, dst in TARGETS:
        path = os.path.join(BASE, src)
        doc = Document(path)
        lines = []
        for p in doc.paragraphs:
            t = clean(p.text)
            if not t:
                continue
            lvl = heading_level(p)
            if lvl:
                lines.append("#" * min(lvl, 6) + " " + t)
            else:
                lines.append(t)
        # 表格转 md
        for ti, table in enumerate(doc.tables):
            if not table.rows:
                continue
            lines.append("")
            lines.append(f"<!-- 表 {ti + 1} -->")
            for row in table.rows:
                cells = [cell_text_marked(c) for c in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")
        out = os.path.join(BASE, dst)
        with open(out, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        print(f"{src} -> {dst} ({len(lines)} lines)")
    run_add_toc()


if __name__ == "__main__":
    main()
